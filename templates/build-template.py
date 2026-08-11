#!/usr/bin/env python3
"""
Transform GIW's real GTP template (source-template.docx) into a
docxtemplater merge-tag template (../backend/src/docx/template.docx).

Re-run this whenever GIW sends an updated source template — see the backend
README for the full workflow. This script is the *only* place placeholder
mapping lives; it doesn't touch the source file, only reads it and writes a
new file.

Why this isn't a simple find-and-replace: Word splits a single visible
placeholder like "[Project Address]" across multiple <w:t> runs whenever a
formatting or spell-check boundary falls inside it, and the cover page date
is a live DATE field (recalculates on open), not plain text. Both need
surgical handling below or a naive regex replace silently fails or, worse,
gets clobbered the next time someone opens the doc in Word.
"""
import copy
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from lxml import etree

SOURCE = Path(__file__).parent / "source-template.docx"
OUTPUT = Path(__file__).parent.parent / "backend" / "src" / "docx" / "template.docx"

# Plain literal-text replacements (no fields involved). Applied paragraph by
# paragraph, across the whole document tree (body, tables, text boxes) plus
# every header/footer part, using a minimal-span run replace so we never
# touch unrelated runs (e.g. an adjacent page-number field) even when the
# placeholder itself is split across runs.
LITERAL_REPLACEMENTS = [
    ("[Project Address]", "{projectAddress}"),
    ("[Client]", "{client}"),
    ("[GIWREF]", "{giwRef}"),
    ("[Council] Planning Scheme", "{councilName} Planning Scheme"),
    ("[Council]", "{councilName}"),
    ("[Architect]", "{architect}"),
    ("[Date]", "{date}"),
    ("Proposed Mixed-Use Development", "Proposed {developmentTypeLabel} Development"),
    (
        "The walkscore of the site is xxx making it suitable to multiple modes of transport.",
        "The indicative walkability score of the site is {indicativeWalkability}/100 (this tool's own "
        "OSM-based proxy, not the commercial Walk Score® product) — confirm with walkscore.com "
        "for a client-facing figure.",
    ),
]

# Whole boilerplate paragraphs replaced by one generated-narrative tag —
# located by a short substring that's unique within the source template and
# survives regardless of processing order (body only; none of these appear
# in headers/footers).
WHOLE_PARAGRAPH_REPLACEMENTS = [
    ("adopt the strategies that promote the usage of public transportation", "{summary}"),
    ("has been engaged by", "{introduction}"),
    ("has an approximate surface area of XXm2", "{subjectSiteNarrative}"),
    ("connects it to the major train lines", "{transportNarrative}"),
    ("Infrastructure Victoria is targeting a zero-emission vehicle infrastructure", "{policyNarrative}"),
    ("is to implement the GTP provided by GIW", "{gtpIntro}"),
    ("requires monitoring and annual review and reporting", "{monitoringAndReporting}"),
]

# (paragraph to convert into a loop template, loop name, per-item tag) —
# the paragraph matched by the substring becomes the single repeatable
# item (keeping its bullet formatting); sibling static paragraphs listed in
# PARAGRAPHS_TO_DELETE are then removed. docxtemplater's paragraphLoop only
# repeats-as-separate-paragraphs when the {#loop} / {/loop} markers are on
# their OWN paragraphs, distinct from the item paragraph — putting all
# three tags in one paragraph (tried first) silently collapses every item
# into one run with no separator between them instead of one bullet per
# item; verified empirically with a minimal render before settling on this.
LOOP_PARAGRAPH_CONVERSIONS = [
    ("Project No. XXX", "sourcesOfInformation", "{.}"),
    ("Reduction of the number of trips by single occupant vehicle", "targets", "{text}"),
]

PARAGRAPHS_TO_DELETE = [
    # 2nd "Sources of Information" bullet is itself split across two <w:p>
    # elements in the source (a manual line break within one bullet) —
    # both need removing. Safe to match "Infrastructure Victoria" here
    # (not just in the Policies paragraph) because WHOLE_PARAGRAPH_REPLACEMENTS
    # runs first and already collapsed that paragraph to "{policyNarrative}".
    "Infrastructure Victoria ",
    "Parliament of Victoria - Report on Climate Change",
    "Planning Scheme",  # 3rd "Sources of Information" bullet, now says "{councilName} Planning Scheme"
    "Use of public transport by at least xx",
    "Use of sustainable modes of transport such as walking, cycling, and shared vehicles by at least xx",
]

# instrText substring -> merge tag to insert in place of the unlinked field.
DATE_FIELDS = [
    (' DATE \\@ "d MMMM yyyy" ', "{generatedDateLong}"),  # cover page
    (' DATE  \\@ "d/MM/yyyy"  \\* MERGEFORMAT', "{generatedDateShort}"),  # "as at" transport intro
]


def iter_runs(paragraph_el):
    # Runs can be nested one level deeper than a direct <w:p> child — inside
    # a content control's <w:sdt>/<w:sdtContent> (this template uses SDTs
    # for several placeholders) or a <w:hyperlink>. iter() finds those;
    # findall() (direct children only) silently misses them.
    return list(paragraph_el.iter(qn("w:r")))


def run_text(run_el):
    return "".join(t.text or "" for t in run_el.findall(qn("w:t")))


def replace_in_paragraph(paragraph_el, old, new):
    """Replace the first occurrence of `old` in a paragraph's visible text
    (concatenated across its runs) with `new`, touching only the minimal
    run span that contains the match. Returns True if a replacement was made."""
    runs = iter_runs(paragraph_el)
    run_texts = [(r, run_text(r)) for r in runs]
    full = "".join(t for _, t in run_texts)
    idx = full.find(old)
    if idx == -1:
        return False
    start, end = idx, idx + len(old)
    pos = 0
    involved = []
    for r, text in run_texts:
        r_start, r_end = pos, pos + len(text)
        if r_end > start and r_start < end:
            involved.append((r, text, r_start, r_end))
        pos = r_end
    if not involved:
        return False
    first_r, first_text, first_start, _ = involved[0]
    last_r, last_text, _, last_end = involved[-1]
    prefix = first_text[: start - first_start]
    last_start = last_end - len(last_text)
    suffix = last_text[end - last_start :]
    combined = prefix + new + suffix

    first_t = first_r.find(qn("w:t"))
    if first_t is None:
        first_t = etree.SubElement(first_r, qn("w:t"))
    first_t.text = combined
    first_t.set(qn("w:space"), "preserve")

    for r, _, _, _ in involved[1:]:
        for t in r.findall(qn("w:t")):
            t.text = ""
    return True


def replace_whole_paragraph(root_el, distinctive_substring, new_tag):
    """Find the paragraph whose text contains `distinctive_substring` and
    replace its ENTIRE text with `new_tag`. Used for boilerplate narrative
    paragraphs (Summary, Introduction, etc.) that are more sensibly
    replaced wholesale by our own generated narrative than surgically
    patched — the original prose has stale example fill-ins ("XX% by
    2020") not worth preserving fragment-by-fragment."""
    for p in root_el.iter(qn("w:p")):
        runs = list(p.iter(qn("w:r")))
        full = "".join(run_text(r) for r in runs)
        if distinctive_substring not in full:
            continue
        if not runs:
            return False
        first_t = runs[0].find(qn("w:t"))
        if first_t is None:
            first_t = etree.SubElement(runs[0], qn("w:t"))
        first_t.text = new_tag
        first_t.set(qn("w:space"), "preserve")
        for r in runs[1:]:
            for t in r.findall(qn("w:t")):
                t.text = ""
        return True
    print(f"  WARNING: whole-paragraph target not found: {distinctive_substring!r}", file=sys.stderr)
    return False


def convert_to_loop_paragraph(root_el, distinctive_substring, loop_name, item_tag):
    """Turn the paragraph matching `distinctive_substring` into the single
    repeatable item of a paragraph-level loop: its own text becomes
    `item_tag` (keeping its formatting, e.g. bullet numPr), with new plain
    marker paragraphs {#loop_name} and {/loop_name} inserted immediately
    before/after it."""
    for p in root_el.iter(qn("w:p")):
        runs = list(p.iter(qn("w:r")))
        full = "".join(run_text(r) for r in runs)
        if distinctive_substring not in full:
            continue
        if not runs:
            return False
        first_t = runs[0].find(qn("w:t"))
        if first_t is None:
            first_t = etree.SubElement(runs[0], qn("w:t"))
        first_t.text = item_tag
        first_t.set(qn("w:space"), "preserve")
        for r in runs[1:]:
            for t in r.findall(qn("w:t")):
                t.text = ""

        def marker_paragraph(text):
            marker_p = etree.Element(qn("w:p"))
            marker_r = etree.SubElement(marker_p, qn("w:r"))
            marker_t = etree.SubElement(marker_r, qn("w:t"))
            marker_t.text = text
            marker_t.set(qn("w:space"), "preserve")
            return marker_p

        p.addprevious(marker_paragraph("{#%s}" % loop_name))
        p.addnext(marker_paragraph("{/%s}" % loop_name))
        return True
    print(f"  WARNING: loop paragraph target not found: {distinctive_substring!r}", file=sys.stderr)
    return False


def delete_paragraph_containing(root_el, distinctive_substring):
    """Delete the paragraph whose text contains `distinctive_substring`.
    Used to remove now-redundant static bullet paragraphs after converting
    one sibling bullet into a docxtemplater loop template."""
    for p in root_el.iter(qn("w:p")):
        runs = list(p.iter(qn("w:r")))
        full = "".join(run_text(r) for r in runs)
        if distinctive_substring in full:
            p.getparent().remove(p)
            return True
    print(f"  WARNING: paragraph to delete not found: {distinctive_substring!r}", file=sys.stderr)
    return False


def replace_everywhere(root_el, old, new):
    count = 0
    for p in root_el.iter(qn("w:p")):
        while replace_in_paragraph(p, old, new):
            count += 1
    return count


def unlink_date_field(root_el, instr_substring, tag):
    """Find the field whose instrText contains `instr_substring`, replace
    the entire fldChar-begin..fldChar-end run sequence with one plain run
    containing `tag`. This is the XML equivalent of Word's Ctrl+Shift+F9
    (unlink field) — necessary so the field can never recalculate over our
    merge tag when someone later opens/prints the generated .docx."""
    for instr in root_el.iter(qn("w:instrText")):
        if instr.text and instr_substring in instr.text:
            run_with_instr = instr.getparent()
            paragraph = run_with_instr.getparent()
            while paragraph is not None and paragraph.tag != qn("w:p"):
                paragraph = paragraph.getparent()
            if paragraph is None:
                continue
            runs = list(paragraph.iter(qn("w:r")))
            try:
                idx = runs.index(run_with_instr)
            except ValueError:
                continue
            begin_idx = None
            for i in range(idx, -1, -1):
                fld = runs[i].find(qn("w:fldChar"))
                if fld is not None and fld.get(qn("w:fldCharType")) == "begin":
                    begin_idx = i
                    break
            end_idx = None
            for i in range(idx, len(runs)):
                fld = runs[i].find(qn("w:fldChar"))
                if fld is not None and fld.get(qn("w:fldCharType")) == "end":
                    end_idx = i
                    break
            if begin_idx is None or end_idx is None:
                print(f"  WARNING: could not find full field bounds for {instr_substring!r}", file=sys.stderr)
                continue
            field_runs = runs[begin_idx : end_idx + 1]
            # Preserve formatting from the run that held the cached display
            # text (usually the one right before fldChar end) if present,
            # else fall back to the begin run's formatting.
            rpr_source = None
            for r in reversed(field_runs):
                if r.findall(qn("w:t")):
                    rpr_source = r
                    break
            if rpr_source is None:
                rpr_source = field_runs[0]
            rpr = rpr_source.find(qn("w:rPr"))

            new_run = etree.Element(qn("w:r"))
            if rpr is not None:
                new_run.append(copy.deepcopy(rpr))
            new_t = etree.SubElement(new_run, qn("w:t"))
            new_t.text = tag
            new_t.set(qn("w:space"), "preserve")

            field_runs[0].addprevious(new_run)
            for r in field_runs:
                r.getparent().remove(r)
            return True
    return False


def process_xml_part(root_el, label, structural=False):
    for instr_substring, tag in DATE_FIELDS:
        if unlink_date_field(root_el, instr_substring, tag):
            print(f"  [{label}] unlinked DATE field -> {tag}")
    for old, new in LITERAL_REPLACEMENTS:
        n = replace_everywhere(root_el, old, new)
        if n:
            print(f"  [{label}] replaced {old!r} -> {new!r} x{n}")

    if not structural:
        return

    for substring, tag in WHOLE_PARAGRAPH_REPLACEMENTS:
        if replace_whole_paragraph(root_el, substring, tag):
            print(f"  [{label}] whole-paragraph -> {tag!r}")

    for substring, loop_name, item_tag in LOOP_PARAGRAPH_CONVERSIONS:
        if convert_to_loop_paragraph(root_el, substring, loop_name, item_tag):
            print(f"  [{label}] loop paragraph -> {loop_name!r} ({item_tag!r})")

    for substring in PARAGRAPHS_TO_DELETE:
        if delete_paragraph_containing(root_el, substring):
            print(f"  [{label}] deleted paragraph containing {substring!r}")


# Action-category tables (indices confirmed by enumerating doc.tables on
# the source file) -> the loop tag name each becomes. Category keys match
# councils.json's actionCategories / actions.json's category field exactly.
ACTION_TABLE_LOOPS = {
    10: "actionsWalking",
    11: "actionsCycling",
    12: "actionsEndOfTripFacilities",
    13: "actionsPublicTransport",
    14: "actionsCarpoolingCarShare",
    15: "actionsCarParking",
    16: "actionsTravelForWorkAmenities",
    17: "actionsManagement",
}


def set_cell_text(cell, new_text):
    p = cell.paragraphs[0]
    p_el = p._p
    runs = list(p_el.iter(qn("w:r")))
    if runs:
        first_t = runs[0].find(qn("w:t"))
        if first_t is None:
            first_t = etree.SubElement(runs[0], qn("w:t"))
        first_t.text = new_text
        first_t.set(qn("w:space"), "preserve")
        for r in runs[1:]:
            for t in r.findall(qn("w:t")):
                t.text = ""
    else:
        p.add_run(new_text)
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def delete_row(table, index):
    row = table.rows[index]
    row._tr.getparent().remove(row._tr)


def process_action_tables(doc):
    for table_idx, loop_tag in ACTION_TABLE_LOOPS.items():
        table = doc.tables[table_idx]
        template_row = table.rows[1]
        set_cell_text(template_row.cells[0], "{#%s}{text}" % loop_tag)
        set_cell_text(template_row.cells[1], "{/%s}" % loop_tag)
        # Remove the other static example rows (2..N) — the loop
        # regenerates one row per item in the council-weighted action list.
        for _ in range(len(table.rows) - 2):
            delete_row(table, 2)
        print(f"  [table {table_idx}] converted to loop '{loop_tag}', {len(table.rows)} rows remain")


def process_transport_table(doc):
    table = doc.tables[5]
    tags = {1: ("{trainStops}", "{trainDistances}"), 2: ("{tramStops}", "{tramDistances}"), 3: ("{busStops}", "{busDistances}")}
    for row_idx, (stops_tag, distances_tag) in tags.items():
        row = table.rows[row_idx]
        set_cell_text(row.cells[1], stops_tag)
        set_cell_text(row.cells[3], distances_tag)
    print("  [table 5] transport summary tags inserted")


def process_revision_table(doc):
    # Confirm the [Date] literal replacement (handled by the generic body
    # pass) actually landed in the revision-history data row.
    table = doc.tables[1]
    text = table.rows[1].cells[1].text
    if "{date}" not in text:
        print(f"  WARNING: revision table date cell is {text!r}, expected '{{date}}'", file=sys.stderr)


def find_paragraph_element(doc, exact_text):
    for p in doc.paragraphs:
        if p.text.strip() == exact_text:
            return p
    raise ValueError(f"paragraph not found: {exact_text!r}")


def set_cell_borders(cell, sz=8, color="00602B"):
    # tcBorders/tcMar must be supplied in schema order (top, left, bottom,
    # right, then insideH/insideV) — see the backend README's docx gotcha
    # note. Not the docx-npm pBdr bug specifically (this is hand-authored
    # XML, not docx-npm output) but the same schema-order requirement
    # applies to any raw OOXML we write by hand.
    tcPr = cell._tc.get_or_add_tcPr()
    borders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for edge in ("top", "left", "bottom", "right"):
        el = etree.SubElement(borders, qn(f"w:{edge}"))
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def insert_draft_banner(doc):
    """Single-cell bordered table (not a paragraph border — see the schema-
    order note above) inserted right before the "Summary" heading, flagging
    every export as an automatically generated draft per the non-functional
    requirement that generated content never look more authoritative than
    it is."""
    anchor = find_paragraph_element(doc, "Summary")
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_borders(cell)
    tcPr = cell._tc.get_or_add_tcPr()
    shading = etree.SubElement(tcPr, qn("w:shd"))
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "FBEAEA")
    set_cell_text(cell, "{draftBanner}")
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    anchor._p.addprevious(table._tbl)
    # add_table() also appends a trailing empty paragraph after the table by
    # default in the body; that's harmless here, leave it.
    print("  inserted draft banner table before 'Summary'")


def insert_car_share_pods_table(doc):
    """New loop-driven table for live nearby car-share pod data, inserted
    right after the branded Flexicar/GoGet/GreenShareCar operator table
    (table 6) and before "Taxi Services" — keeps the branded operator
    callouts from the original template intact while adding the more
    useful live-data table."""
    anchor = find_paragraph_element(doc, "Taxi Services")
    heading = doc.add_paragraph("Nearby Car Share Pods (live data)", style="Heading 3")
    intro = doc.add_paragraph(
        "{#hasCarSharePods}The following car-share pods were identified within the search radius via "
        "OpenStreetMap:{/hasCarSharePods}{^hasCarSharePods}No car-share pods were identified within the "
        "search radius via OpenStreetMap — confirm manually.{/hasCarSharePods}"
    )
    table = doc.add_table(rows=2, cols=3)
    table.style = doc.tables[5].style
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, ["Name", "Distance", "Walk time"]):
        set_cell_text(cell, text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    row_cells = table.rows[1].cells
    set_cell_text(row_cells[0], "{#carSharePods}{name}")
    set_cell_text(row_cells[1], "{distanceLabel}")
    set_cell_text(row_cells[2], "{walkMinutes} min{/carSharePods}")

    anchor._p.addprevious(heading._p)
    heading._p.addnext(intro._p)
    intro._p.addnext(table._tbl)
    print("  inserted car-share pods loop table before 'Taxi Services'")


def main():
    if not SOURCE.exists():
        sys.exit(f"Source template not found: {SOURCE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document(str(SOURCE))

    print("Body:")
    process_xml_part(doc.element.body, "body", structural=True)

    print("\nTables:")
    process_action_tables(doc)
    process_transport_table(doc)
    process_revision_table(doc)

    print("\nStructural insertions:")
    insert_draft_banner(doc)
    insert_car_share_pods_table(doc)

    for i, section in enumerate(doc.sections):
        for kind, part in [("header", section.header), ("footer", section.footer),
                            ("first_page_header", section.first_page_header),
                            ("first_page_footer", section.first_page_footer)]:
            if part is not None and part.is_linked_to_previous is False:
                process_xml_part(part._element, f"section{i}.{kind}")

    doc.save(str(OUTPUT))
    print(f"\nWrote {OUTPUT}")


if __name__ == "__main__":
    main()
